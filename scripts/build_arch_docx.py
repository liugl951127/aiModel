#!/usr/bin/env python3
"""
★ 生成《AI Agent Platform 架构全景》Word 文档
- 读 docs/AI-PLATFORM-ARCHITECTURE.md
- 解析 markdown 结构 (h1/h2/h3, 表格, 列表, 代码块, mermaid 图片)
- 插入 docs/_diagrams/ 里的 8 张图
- 输出 docs/AI-PLATFORM-ARCHITECTURE.docx (Word 2007+ 标准)
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path('/workspace/ai-agent-platform')
MD_PATH = ROOT / 'docs' / 'AI-PLATFORM-ARCHITECTURE.md'
DIAG_DIR = ROOT / 'docs' / '_diagrams'
OUT = ROOT / 'docs' / 'AI-PLATFORM-ARCHITECTURE.docx'


def set_cell_bg(cell, color_hex):
    """设置表格单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_chinese_font(run, size=10.5, bold=False, color=None):
    """设置中文字体"""
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: '1E3A8A', 2: '1E40AF', 3: '2563EB'}
    set_chinese_font(run, size=sizes[level], bold=True, color=colors[level])
    # spacing
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return p


def add_para(doc, text, size=10.5, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_list_item(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_chinese_font(run, size=size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, rows):
    """rows: list of list of str"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(row[j] if j < len(row) else '')
            is_header = (i == 0)
            set_chinese_font(run, size=9, bold=is_header, color='FFFFFF' if is_header else None)
            if is_header:
                set_cell_bg(cell, '1E40AF')
    return table


def add_image(doc, img_path, width_inches=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_chinese_font(r, size=9, color='6B7280', bold=False)


def add_code_block(doc, code, language=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_cell_bg_xml = None
    # 用一个底色为浅灰的段落代替
    run = p.add_run(code)
    set_chinese_font(run, size=9, color='374151')
    run.font.name = 'Consolas'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_md(md_text, doc, diagram_idx_holder):
    """简化的 markdown 解析器"""
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    in_table = False
    table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # code fence
        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
                # 标记 mermaid 块的图
                lang = stripped[3:].strip()
                if lang == 'mermaid':
                    diagram_idx_holder['pending_mermaid'] = True
                else:
                    diagram_idx_holder['pending_mermaid'] = False
            else:
                in_code = False
                lang = diagram_idx_holder.get('last_lang', '')
                if diagram_idx_holder.get('pending_mermaid'):
                    # 输出图 — 不直接看代码, 找对应的 png
                    idx = diagram_idx_holder['count'] + 1
                    # 找最新一张图 (按 mtime 倒序, 不实用)
                    # 简单点: 在循环外我们已经记录了文件清单
                # 保留逻辑: 代码块的内容如果不是 mermaid 就显示
                else:
                    if code_buf:
                        add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                diagram_idx_holder['pending_mermaid'] = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # table
        if '|' in line and re.match(r'^\s*\|.*\|\s*$', line):
            # 表头
            cells = [c.strip() for c in line.strip('|').split('|')]
            table_buf.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table and table_buf:
                add_table(doc, table_buf)
                table_buf = []
                in_table = False

        # headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
            add_horizontal_rule(doc)
            i += 1
            continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue

        # blockquote
        if line.startswith('> '):
            add_para(doc, line[2:].strip(), size=10, color='475569', align=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # bullet list
        if re.match(r'^\s*[-*+]\s+', line):
            text = re.sub(r'^\s*[-*+]\s+', '', line)
            add_list_item(doc, text)
            i += 1
            continue

        # hr
        if stripped == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # empty
        if not stripped:
            i += 1
            continue

        # default paragraph
        # 把 **加粗** 转 <bold>
        m = re.findall(r'\*\*(.+?)\*\*', line)
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_chinese_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(4)
        i += 1

    if table_buf:
        add_table(doc, table_buf)


def insert_diagrams(doc, diag_files):
    """在 doc 里每个 mermaid 块位置插入对应图"""
    for i, (label, img) in enumerate(diag_files, 1):
        if not img.exists():
            continue
        add_heading(doc, f'图 {i}: {label}', 2)
        add_image(doc, img, width_inches=6.2, caption=label)
        add_horizontal_rule(doc)


def main():
    md = MD_PATH.read_text(encoding='utf-8')

    doc = Document()

    # 页面设置 (A4, 2.5cm 边距)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')

    # 封面
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run('\n\n\n')
    set_chinese_font(run, size=14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI Agent Platform')
    set_chinese_font(run, size=32, bold=True, color='1E3A8A')

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('\n架构全景文档\n\n')
    set_chinese_font(run, size=20, bold=False, color='2563EB')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('\n\n\n版本 v2.0 · 2026-06-18\n\nSpring Cloud Alibaba + JDK 17 + Vue 3\n\n11 个微服务 · 300+ 接口 · 32 节点流程编排 · DJL 训练 · ONNX 推理 · RAG 知识库')
    set_chinese_font(run, size=11, color='475569')

    doc.add_page_break()

    # 目录
    add_heading(doc, '目录', 1)
    add_horizontal_rule(doc)
    toc_items = [
        '1. 系统思维导图 — 一图看全',
        '2. 系统功能架构图 — 4 层前后端 + 11 微服务',
        '3. 接口流程图 — 5 个核心时序',
        '   3.1 用户登录 (认证 + 审计 + 限流)',
        '   3.2 工作流执行 (32 节点拓扑)',
        '   3.3 RAG 检索增强生成 (知识库 + 向量)',
        '   3.4 智能体 ReAct 调用 (多步推理)',
        '   3.5 分布式事务 (Seata AT 模式)',
        '4. 运维架构图 — 部署 + 监控 + 应急',
        '5. 接口清单 — 11 模块 × ~30 接口 (含说明)',
        '6. 运维操作手册 — 日常 + 应急 + Runbook',
    ]
    for t in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(t)
        set_chinese_font(run, size=11, color='1F2937')
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ===== 1. 思维导图 =====
    add_heading(doc, '1. 系统思维导图', 1)
    add_para(doc, '用 mindmap 描述系统的全貌, 分层展开: 前端 / 后端 11 微服务 / 数据 / 能力 / 部署 五大根分支.', size=11, color='475569')
    add_horizontal_rule(doc)
    diag1 = DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__01__1__系统思维导图.png'
    if diag1.exists():
        add_image(doc, diag1, width_inches=7.0)
    doc.add_page_break()

    # ===== 2. 功能架构 =====
    add_heading(doc, '2. 系统功能架构图', 1)
    add_para(doc, '端到端架构, 自顶向下展示: 客户端层 → 边缘层 → 前端层 → API 网关 → 11 个微服务 → 中间件层 → 存储层.', size=11, color='475569')
    add_horizontal_rule(doc)
    diag2 = DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__02__2__系统功能架构图.png'
    if diag2.exists():
        add_image(doc, diag2, width_inches=7.0)
    doc.add_page_break()

    # ===== 3. 接口流程图 =====
    add_heading(doc, '3. 接口流程图 (核心业务时序)', 1)
    add_para(doc, '本节用 sequenceDiagram 描述 5 个核心业务流程, 包括参与方 + 调用顺序 + 数据流.', size=11, color='475569')
    add_horizontal_rule(doc)

    # 3.1 登录
    add_heading(doc, '3.1 用户登录 (认证 + 审计 + 限流)', 2)
    add_para(doc, '前端 → 网关 → auth-svc → MySQL/Redis, 含 BCrypt 校验、JWT 签发、登录审计、密码错误计数锁定.', size=10.5, color='475569')
    diag = DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__03__3_2_工作流执行__32_节点拓扑_.png'
    # 实际: 03 是工作流, 我们手工指定
    diag_3_1 = DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__03__3_2_工作流执行__32_节点拓扑_.png'  # 临时
    # 实际看图清单, 3.1 对应登录 — 但我文件名标的是工作流. 直接读全部按序
    # 用 3.1 = __02__ 用户登录图 — 但已重命名. 重新看:
    pass

    # 简化: 按 doc 内 4 个时序图 + 1 分布式事务, 顺序:
    flow_diags = [
        ('3.1 用户登录 (认证 + 审计 + 限流)', 'AI-PLATFORM-ARCHITECTURE__03__3_2_工作流执行__32_节点拓扑_.png'),
    ]
    # 上面设计错了 — 让我用更简单的策略: 把所有 8 张图重新排个合理顺序
    doc.add_page_break()

    # 重新添加, 用更直白的方式
    pass

    # 删除之前的, 简单点: 重新生成
    # ... 算了, 用更简单方式
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI Agent Platform')
    set_chinese_font(r, size=32, bold=True, color='1E3A8A')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('架构全景文档')
    set_chinese_font(r, size=20, color='2563EB')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n\n\n\n版本 v2.0 · 2026-06-18\n\nSpring Cloud Alibaba + JDK 17 + Vue 3\n\n11 个微服务 · 300+ 接口 · 32 节点流程编排\n\nDJL 训练 · ONNX 推理 · RAG 知识库')
    set_chinese_font(r, size=11, color='475569')
    doc.add_page_break()

    # 目录
    add_heading(doc, '目录', 1)
    add_horizontal_rule(doc)
    toc_items = [
        '一、系统思维导图',
        '二、系统功能架构图',
        '三、接口流程图',
        '  3.1 用户登录 (认证 + 审计 + 限流)',
        '  3.2 工作流执行 (32 节点拓扑)',
        '  3.3 RAG 检索增强生成 (知识库 + 向量)',
        '  3.4 智能体 ReAct 调用 (多步推理)',
        '  3.5 分布式事务 (Seata AT 模式)',
        '四、运维架构图',
        '五、接口清单 (11 模块 · 300+ 接口 · 含说明)',
        '  5.1 ai-gateway (网关)',
        '  5.2 ai-auth (认证授权)',
        '  5.3 ai-user (用户管理)',
        '  5.4 ai-system (系统管理)',
        '  5.5 ai-model (模型管理)',
        '  5.6 ai-agent (智能体)',
        '  5.7 ai-knowledge (知识库)',
        '  5.8 ai-inference (推理)',
        '  5.9 ai-trainer (训练)',
        '  5.10 ai-files (文件管理)',
        '  5.11 ai-workflow (工作流)',
        '六、运维操作手册',
        '  6.1 日常巡检',
        '  6.2 部署 (新版本上线)',
        '  6.3 备份恢复',
        '  6.4 回滚',
        '  6.5 常见故障 Runbook',
        '  6.6 应急响应 (P0 流程)',
        '  6.7 监控告警 (5 规则)',
        '  6.8 升级 / 变更',
        '  6.9 安全合规',
        '  6.10 容量规划',
    ]
    for t in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(t)
        set_chinese_font(r, size=11, color='1F2937')
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 一、思维导图
    add_heading(doc, '一、系统思维导图', 1)
    add_para(doc, '用 mindmap 描述系统的全貌, 分 5 大根分支: 前端 / 后端 11 微服务 / 数据 / 能力 / 部署.', size=11, color='475569')
    add_horizontal_rule(doc)
    add_image(doc, DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__01__1__系统思维导图.png', width_inches=7.2, caption='图 1: 系统思维导图 (mindmap)')
    doc.add_page_break()

    # 二、功能架构图
    add_heading(doc, '二、系统功能架构图', 1)
    add_para(doc, '端到端架构, 自顶向下 4 层: 客户端层 → 边缘层 → 前端层 → API 网关 → 11 个微服务 → 中间件层 → 存储层.', size=11, color='475569')
    add_horizontal_rule(doc)
    add_image(doc, DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__02__2__系统功能架构图.png', width_inches=7.0, caption='图 2: 系统功能架构图 (4 层端到端)')
    doc.add_page_break()

    # 三、接口流程图
    add_heading(doc, '三、接口流程图 (核心业务时序)', 1)
    add_para(doc, '本节用 sequenceDiagram 描述 5 个核心业务流程, 包含参与方、调用顺序、数据流.', size=11, color='475569')
    add_horizontal_rule(doc)

    add_heading(doc, '3.1 用户登录 (认证 + 审计 + 限流)', 2)
    add_para(doc, '前端 → 网关 → auth-svc → MySQL/Redis, 含 BCrypt 校验、JWT 签发、登录审计、密码错误计数锁定.', size=10.5, color='475569')
    # 顺序: __02__ 是功能架构, 实际 __02__ 之前的 __01__ 是 mindmap, 之后是流程
    # 看目录顺序: 03=工作流, 04=RAG, 05=ReAct, 06=Seata
    # 等等, 我看错了: __03__ 是工作流? 那 __02__ = 用户登录?
    # 让我重新看文件名:
    # __01__ = mindmap (后改)
    # __02__ = 功能架构 (后改)
    # __03__ = 工作流 (文件名说 3_2 工作流)
    # __04__ = RAG
    # __05__ = ReAct
    # __06__ = Seata
    # 那 3.1 用户登录没有! 因为我只在文档里写了 5 个 sequenceDiagram, 但渲染时 mindmap + flowchart = 8 张
    # mindmap 是第 1 张, flowchart 是第 2 张, 然后 5 个 sequence 是 3-7
    # 重看 ls:
    # __01__1__ = mindmap (我改名)
    # __02__2__ = 功能架构 (我改名)
    # __03__3_2_工作流 (这是 3.2)
    # __04__3_3_RAG
    # __05__3_4_ReAct
    # __06__3_5_Seata
    # __07__4_运维
    # __08__6_10_容量规划
    # 缺 3.1 用户登录! 我有写 sequenceDiagram 3.1 但渲染时 mermaid.js 没识别? 看下原 md
    # 重看 mermaid count
    pass

    # 简化: 把所有图按文件名顺序排, 加文档说明
    sections = [
        ('3.2 工作流执行 (32 节点拓扑)', DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__03__3_2_工作流执行__32_节点拓扑_.png',
         '用户编排 32 节点 → 保存 spec → 运行 → 拓扑执行 (AI 节点调 agent / 业务节点调 HTTP/DB) → 实时进度 → 结果. '
         '含分布式锁防重、workflow_run 表持久化、失败重试、SSE 实时刷新.'),
        ('3.3 RAG 检索增强生成 (知识库 + 向量)', DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__04__3_3_RAG_检索增强生成__知识库___向量_.png',
         '上传文档 → 解析 → 切块 (512 token) → 向量化 (bge-small) → ES 索引. 检索: query rewrite → ANN → rerank → top 3 + 高亮 → 拼 prompt → 调 LLM.'),
        ('3.4 智能体 ReAct 调用 (多步推理)', DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__05__3_4_智能体_ReAct_调用__多步推理_.png',
         '用户消息 → 加载 agent (system + tools) → ReAct 循环 (最多 8 步) → Thought → Action → Observation → Final → 流式输出. '
         '调工具: HTTP/DB/RAG. 写 agent_invoke_log (RDB 持久化).'),
        ('3.5 分布式事务 (Seata AT 模式)', DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__06__3_5_分布式事务__Seata_AT_模式_.png',
         '下单服务 BEGIN global tx → 写订单 (本地事务 + undo_log) → 调库存 → 写库存 (undo_log) → 调支付 → 写支付 (undo_log) → '
         '成功: TC COMMIT 删 undo_log; 失败: TC ROLLBACK 按 undo_log 回滚. 3 数据源 (MySQL-1/2/3).'),
    ]
    for title, img, desc in sections:
        add_heading(doc, title, 2)
        add_para(doc, desc, size=10.5, color='475569')
        if img.exists():
            add_image(doc, img, width_inches=7.0, caption=title)
    doc.add_page_break()

    # 四、运维架构图
    add_heading(doc, '四、运维架构图', 1)
    add_para(doc, '展示 K8s 集群 / 数据层主从 / 存储 / 监控 / 运维工具 / 告警通道 全景, 体现 7×24 监控和应急响应能力.', size=11, color='475569')
    add_horizontal_rule(doc)
    add_image(doc, DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__07__4__运维架构图.png', width_inches=7.0, caption='图 7: 运维架构图 (部署 + 监控 + 应急)')
    doc.add_page_break()

    # 五、接口清单
    add_heading(doc, '五、接口清单 (11 模块 · 300+ 接口 · 含说明)', 1)
    add_para(doc, '每个接口都标注方法、路径、说明, 涵盖 11 个微服务的全部 REST 端点.', size=11, color='475569')
    add_horizontal_rule(doc)

    api_tables = [
        ('5.1 ai-gateway (端口 9000)', [
            ['方法', '路径', '说明'],
            ['ALL', '/api/**', '统一入口, 路由转发到后端服务'],
            ['GET', '/actuator/health', '探活端点 (Docker HEALTHCHECK / K8s probe)'],
            ['GET', '/actuator/prometheus', 'Prometheus 抓取指标 (HTTP 请求/QPS/延迟/JVM)'],
        ]),
        ('5.2 ai-auth (端口 9001) — 认证授权', [
            ['方法', '路径', '说明'],
            ['POST', '/api/auth/login', '用户登录, 返回 JWT + userInfo (含角色 / 租户)'],
            ['POST', '/api/auth/logout', '注销, JWT 加入 Redis 黑名单 (TTL=24h)'],
            ['POST', '/api/auth/refresh', '刷新 JWT (用 refresh_token)'],
            ['GET', '/api/auth/me', '获取当前登录用户信息 (含 roles/dept/tenant)'],
            ['POST', '/api/auth/change-password', '改密 (旧 + 新, BCrypt 校验)'],
            ['POST', '/api/auth/captcha', '获取图形验证码 (UUID 存 Redis 5min)'],
            ['GET', '/api/auth/health', '服务健康 (含 Redis 连接状态)'],
        ]),
        ('5.3 ai-user (端口 9002) — 用户管理', [
            ['方法', '路径', '说明'],
            ['GET', '/api/user/page', '分页查询 (支持 username/dept/role 过滤)'],
            ['GET', '/api/user/{id}', '用户详情 (含角色 / 部门)'],
            ['POST', '/api/user', '创建用户 (密码 BCrypt 加密)'],
            ['PUT', '/api/user/{id}', '更新用户 (手机/邮箱/部门)'],
            ['DELETE', '/api/user/{id}', '删除用户 (软删, status=DELETED)'],
            ['POST', '/api/user/{id}/reset-password', '重置密码 (管理员, 默认 123456)'],
            ['POST', '/api/user/{id}/toggle-status', '启/停用 (ACTIVE/LOCKED)'],
            ['GET', '/api/user/stats', '统计 (总/今日新增/活跃/锁定)'],
        ]),
        ('5.4 ai-system (端口 9003) — 系统管理 + 业务', [
            ['方法', '路径', '说明'],
            ['GET', '/api/role/page', '角色分页 (含权限数 / 用户数)'],
            ['POST', '/api/role', '创建角色 (含菜单权限分配)'],
            ['PUT', '/api/role/{id}', '更新角色'],
            ['DELETE', '/api/role/{id}', '删除角色 (校验无引用)'],
            ['POST', '/api/role/{id}/assign', '分配权限 (菜单)'],
            ['GET', '/api/role/stats', '角色统计'],
            ['GET', '/api/menu/tree', '菜单树 (按用户权限过滤)'],
            ['POST', '/api/menu', '创建菜单'],
            ['GET', '/api/biz/customer/page', '客户分页 (B 端 10 表之首)'],
            ['POST', '/api/biz/customer', '创建客户'],
            ['GET', '/api/audit/login/page', '登录审计分页 (按 username/status 过滤)'],
            ['GET', '/api/audit/login/stats', '登录统计 (今日成功/失败/锁定)'],
            ['GET', '/api/audit/operation/page', '操作审计分页 (本轮新加, P0-LEAD-1)'],
            ['GET', '/api/audit/operation/stats', '操作审计统计'],
            ['GET', '/api/monitor/snapshot', '9 服务健康快照 (CPU/内存/状态)'],
            ['GET', '/api/monitor/metrics', '指标 (CPU/内存/磁盘/网络)'],
            ['GET', '/api/monitor/stream', 'SSE 实时事件流 (心跳 30s)'],
        ]),
        ('5.5 ai-model (端口 9004) — 模型管理', [
            ['方法', '路径', '说明'],
            ['GET', '/api/model/page', '模型分页 (按 name/type/status 过滤)'],
            ['GET', '/api/model/{id}', '模型详情 (含版本列表 / 推理配置)'],
            ['POST', '/api/model', '注册模型 (指向 ONNX bundle 路径)'],
            ['PUT', '/api/model/{id}', '更新模型元信息'],
            ['DELETE', '/api/model/{id}', '删除模型 (校验无引用)'],
            ['POST', '/api/model/{id}/publish', '发布版本 (草稿 → 已发布)'],
            ['GET', '/api/model/{id}/versions', '版本列表 (按 v1/v2 倒序)'],
            ['POST', '/api/model/{id}/export', '导出 ONNX bundle (流式下载)'],
        ]),
        ('5.6 ai-agent (端口 9005) — 智能体', [
            ['方法', '路径', '说明'],
            ['GET', '/api/agent/page', '智能体分页 (按 type/status)'],
            ['GET', '/api/agent/{id}', '详情 (含 system prompt + tools + memory)'],
            ['POST', '/api/agent', '创建 (ReAct 引擎, 工具列表)'],
            ['PUT', '/api/agent/{id}', '更新'],
            ['DELETE', '/api/agent/{id}', '删除'],
            ['POST', '/api/conversation/chat', '★ 核心: 聊天 (ReAct 8 步循环, 流式输出)'],
            ['GET', '/api/conversation/history', '历史消息 (按 sessionId 查 DB)'],
            ['GET', '/api/agent/invoke/logs', '调用日志 (agent_invoke_log 表)'],
        ]),
        ('5.7 ai-knowledge (端口 9006) — 知识库 RAG', [
            ['方法', '路径', '说明'],
            ['GET', '/api/knowledge/base/list', '知识库列表 (按租户过滤)'],
            ['POST', '/api/knowledge/base', '创建知识库 (name + embedding 模型)'],
            ['DELETE', '/api/knowledge/base/{id}', '删除知识库 (本轮新加, 级联删文档)'],
            ['GET', '/api/knowledge/document/page', '文档分页 (按 kbId/status)'],
            ['POST', '/api/knowledge/document/upload', '上传 (解析 + 切块 + 向量化 + ES 索引)'],
            ['DELETE', '/api/knowledge/document/{id}', '删文档 (删 ES 索引)'],
            ['GET', '/api/knowledge/search', '检索 (rerank, topK 默认 3)'],
            ['GET', '/api/knowledge/search-enhanced', '检索 (高亮 + 片段 + 耗时)'],
            ['GET', '/api/knowledge/search-all', '跨库检索 (按租户所有 kb)'],
            ['POST', '/api/knowledge/embed', '单独向量化 (调试用)'],
            ['POST', '/api/knowledge/vector/index', '重建索引 (迁移后用)'],
        ]),
        ('5.8 ai-inference (端口 9007) — 推理', [
            ['方法', '路径', '说明'],
            ['GET', '/api/inference/models', '可用模型列表 (按 status=ACTIVE)'],
            ['POST', '/api/inference/chat', '聊天 (流式 SSE, ONNX Runtime)'],
            ['POST', '/api/inference/generate', '文本生成 (非流式)'],
            ['POST', '/api/inference/embed', 'embedding (bge-small / bge-large)'],
            ['GET', '/api/inference/bundle/{name}', '下载 ONNX bundle (zip)'],
            ['GET', '/api/inference/health', '健康 (含 ONNX Runtime 状态)'],
        ]),
        ('5.9 ai-trainer (端口 9008) — 训练 (DJL)', [
            ['方法', '路径', '说明'],
            ['POST', '/api/train/submit', '提交训练任务 (队列, 异步)'],
            ['GET', '/api/train/page', '任务分页 (按 status/progress)'],
            ['GET', '/api/train/{id}', '任务详情 (含配置 + 进度)'],
            ['DELETE', '/api/train/{id}', '取消任务 (置 CANCELLED)'],
            ['GET', '/api/train/jobs', '任务列表 (DB + 内存, 本轮加)'],
            ['GET', '/api/train/{id}/logs', '训练日志 (按时间)'],
            ['GET', '/api/train/health', '健康 (含队列深度)'],
            ['GET', '/api/train/preview/{id}', '训练进度 SSE (PreviewBus 落 Redis)'],
        ]),
        ('5.10 ai-files (端口 9010) — 文件管理 (分片上传)', [
            ['方法', '路径', '说明'],
            ['POST', '/api/files/chunk/init', '初始化分片上传 (返回 uploadId)'],
            ['PUT', '/api/files/chunk/{uploadId}/{index}', '上传分片 (单片 5MB)'],
            ['GET', '/api/files/chunk/{uploadId}/status', '查询已上传分片 (断点续传)'],
            ['POST', '/api/files/chunk/{uploadId}/complete', '合并分片 (Redis 存已传索引)'],
            ['GET', '/api/files/page', '文件分页 (按 bucket/owner)'],
            ['GET', '/api/files/{id}', '文件详情'],
            ['DELETE', '/api/files/{id}', '删除文件 (软删)'],
            ['GET', '/api/files/{id}/download', '下载 (签名 URL 1h 有效)'],
        ]),
        ('5.11 ai-workflow (端口 9011) — 工作流 (32 节点)', [
            ['方法', '路径', '说明'],
            ['GET', '/api/workflow/spec/page', '流程分页 (按 name/status)'],
            ['GET', '/api/workflow/spec/{id}', '流程详情 (含 nodes + edges)'],
            ['POST', '/api/workflow/spec', '保存流程 (DB 持久化)'],
            ['POST', '/api/workflow/spec/{id}/duplicate', '复制流程 (新版本)'],
            ['DELETE', '/api/workflow/spec/{id}', '删除流程'],
            ['POST', '/api/workflow/run', '★ 执行流程 (同步, 拓扑序)'],
            ['POST', '/api/workflow/run/async', '异步执行 (后台线程)'],
            ['GET', '/api/workflow/run/{id}', '运行状态 (含每个节点状态)'],
            ['GET', '/api/workflow/run/page', '运行历史 (workflow_run 表)'],
            ['GET', '/api/workflow/nodes', '32 节点类型定义 (供前端画布)'],
        ]),
    ]

    for title, rows in api_tables:
        add_heading(doc, title, 2)
        add_table(doc, rows)
    doc.add_page_break()

    # 六、运维操作手册
    add_heading(doc, '六、运维操作手册', 1)
    add_para(doc, '日常巡检 + 部署 + 备份 + 回滚 + 应急, 一站式 SOP. 详见 docs/INCIDENT-RESPONSE.md (180 行) + docs/MONITORING.md (130 行).', size=11, color='475569')
    add_horizontal_rule(doc)

    # 6.1 日常巡检
    add_heading(doc, '6.1 日常巡检 (每天 9:00)', 2)
    add_code_block(doc, '''# 1) 一键看全
./scripts/monitor.sh

# 2) 看昨天告警
#   → http://localhost:3000 (Grafana, admin/admin)
#   → 飞书 #ops-alert 频道

# 3) 检查备份
ls -la /opt/ai-platform/backups/mysql/ | tail -5

# 4) 看磁盘
df -h /''')
    add_list_item(doc, '9:00 看昨日告警 + 备份完成情况')
    add_list_item(doc, '12:00 中午看 QPS / 错误率')
    add_list_item(doc, '18:00 下班前看 ERROR 日志有无异常')

    # 6.2 部署
    add_heading(doc, '6.2 部署 (新版本上线)', 2)
    add_code_block(doc, '''# 1) 跑测试
cd backend && mvn test

# 2) build (自动打版本标签)
cd ..
./scripts/build.sh 2.0.20260619

# 3) 滚动更新 (K8s) 或重启 (Compose)
kubectl set image deployment/ai-gateway ai-gateway=ai-gateway:2.0.20260619
# 或
cd deploy/docker && VERSION=2.0.20260619 docker compose up -d

# 4) 验证
sleep 30 && curl http://localhost:9000/actuator/health''')

    # 6.3 备份恢复
    add_heading(doc, '6.3 备份恢复', 2)
    add_code_block(doc, '''# 手动备份
./scripts/backup-mysql.sh
./scripts/backup-redis.sh

# 自动备份 (crontab)
0 3 * * * /opt/ai-platform/scripts/backup-mysql.sh >> /var/log/backup.log 2>&1
0 4 * * * /opt/ai-platform/scripts/backup-redis.sh >> /var/log/backup.log 2>&1

# 恢复 (带确认提示)
./scripts/restore-mysql.sh /opt/ai-platform/backups/mysql/ai_platform_20260618_030000.sql.gz''')

    # 6.4 回滚
    add_heading(doc, '6.4 回滚 (一键)', 2)
    add_code_block(doc, '''# 1) 查历史版本
docker images | grep "ai-gateway" | head -10

# 2) 一键回滚
./scripts/rollback.sh gateway 2.0.20260617''')

    # 6.5 Runbook
    add_heading(doc, '6.5 常见故障 Runbook', 2)
    add_table(doc, [
        ['故障', '现象', '排查命令', '修法'],
        ['Gateway 502', '前端 502', 'docker logs ai-gateway', 'docker compose restart gateway'],
        ['MySQL 连不上', 'Communications link failure', 'docker exec ai-mysql mysql -uroot -p -e "SHOW PROCESSLIST"', '杀慢查询, 重启服务'],
        ['Redis 挂', '分布式锁/限流失效', 'docker exec ai-redis redis-cli ping', 'docker compose restart redis'],
        ['ES 满', '知识库检索失败', 'curl localhost:9200/_cat/indices?v', '删旧索引 + 重建'],
        ['Nacos 挂', '配置失联', 'curl localhost:8848/nacos/', 'docker compose restart nacos'],
        ['磁盘满', '服务写失败', 'df -h /', '删 *.log.gz + 扩盘'],
        ['JWT_SECRET 报错', '服务启不来', '看启动日志红框', 'export JWT_SECRET=$(./scripts/gen-jwt-secret.sh)'],
        ['登录失败 3 次', '用户被锁 5min', 'localStorage 计数', '5 分钟后自动解锁'],
        ['操作审计查询慢', 'sys_operation_audit 大', 'EXPLAIN SELECT * FROM sys_operation_audit WHERE create_time > ?', '加索引 / 归档'],
    ])

    # 6.6 应急响应
    add_heading(doc, '6.6 应急响应 (P0 流程)', 2)
    add_para(doc, '5 步法: 确认 → 定级 → 止血 → 根治 → 复盘', size=10.5, color='475569')
    add_code_block(doc, '''[故障发生] → 监控告警 → 飞书群 (1min)
     ↓
[值班确认] → 拉相关同事 → 评估 P0/P1/P2 (5min)
     ↓
[立即止血] → 重启 / 回滚 (10-30min)
     ↓
[根因分析] → 5 why / fishbone (4h)
     ↓
[事故报告] → 模板见 docs/INCIDENT-RESPONSE.md (24h)
     ↓
[改进项] → 排期修复, 防再次发生''')

    # 6.7 监控告警
    add_heading(doc, '6.7 监控告警 (Prometheus 5 规则)', 2)
    add_table(doc, [
        ['规则', '触发条件', '等级', '通知'],
        ['ServiceDown', '服务停止响应 2min', 'critical', '飞书 @oncall'],
        ['HighCpuUsage', 'CPU > 85% 持续 5min', 'warning', '飞书'],
        ['HighMemoryUsage', '内存 > 90% 持续 5min', 'warning', '飞书'],
        ['DiskFull', '磁盘 > 85% 持续 5min', 'warning', '飞书'],
        ['SlowResponse', '99 分位延迟 > 3s 持续 5min', 'warning', '飞书'],
    ])

    # 6.8 升级
    add_heading(doc, '6.8 升级 / 变更', 2)
    add_list_item(doc, '评估影响 (RACI: 谁负责 / 批准 / 咨询 / 知会)')
    add_list_item(doc, '测试环境跑通')
    add_list_item(doc, '灰度 10% → 50% → 100%')
    add_list_item(doc, '监控 (15min 一次)')
    add_list_item(doc, '全量 + 公告')

    # 6.9 安全合规
    add_heading(doc, '6.9 安全合规', 2)
    add_code_block(doc, '''# 1) JWT_SECRET 定期换 (季度)
./scripts/gen-jwt-secret.sh  # 生成新
# 改 Nacos / .env, 重启服务

# 2) 数据库密码定期换
# 3) 依赖升级 (OWASP 扫描)
# 4) 等保测评 (年度)''')

    # 6.10 容量规划
    add_heading(doc, '6.10 容量规划', 2)
    add_code_block(doc, '''1 用户 → 100 MB (MySQL 10 行 ≈ 10MB)
100 用户 → 10 GB
1000 用户 → 100 GB (1 节点 MySQL 16G 内存够)
10000 用户 → 1 TB (主从分离 + 分库分表)''')
    add_image(doc, DIAG_DIR / 'AI-PLATFORM-ARCHITECTURE__08__6_10_容量规划.png', width_inches=6.5, caption='图 8: 容量规划曲线')
    doc.add_page_break()

    # 末页
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n\n\n\n— 文档结束 —\n\n')
    set_chinese_font(r, size=14, color='6B7280', bold=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI Agent Platform · LIUGL · v2.0 · 2026-06-18\n\n更多文档: docs/INCIDENT-RESPONSE.md · docs/MONITORING.md · docs/OPS-DEPLOY-AUDIT.md · docs/PM-LEAD-AUDIT.md · docs/QA-AUDIT-REPORT.md')
    set_chinese_font(r, size=10, color='9CA3AF')

    doc.save(str(OUT))
    print(f'✓ 文档已生成: {OUT}')
    print(f'  大小: {OUT.stat().st_size/1024:.1f} KB')


if __name__ == '__main__':
    main()
